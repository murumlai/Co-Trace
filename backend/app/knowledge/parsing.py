"""Discovery, categorization, and parsing of supporting product documents.

Filename conventions (verified against real sample docs):

* ``M79060-001_Debug Support Document Key Learning.pdf`` -> ``debug_learning``
* ``M13983-700_Sedona USB Test Card.pdf``               -> ``product_overview``
* ``N32828-201_MPDU_20V_PDB1_HLD_Rev05.docx``           -> ``hld``
* ``N32828_RFC.xlsx``                                   -> ``rfc_knowledge``
* ``N32828-201_RFC.xlsx``                               -> ``rfc_knowledge``

Product code = the first filename token when it looks like a product code,
otherwise the first product-code match anywhere in the filename.

Product-family code = the base code (letters + digits, no revision suffix)
extracted from the filename; used for family-level fallback when a revisioned
product has no exact RFC workbook.

PDF, DOCX, and XLSX (RFC workbooks) are supported. PPTX/legacy Office formats
are planned parser-adapter extension points and intentionally not implemented.
"""
from __future__ import annotations

import glob
import hashlib
import os
import re

from ..config import settings
from .models import DocumentCategory, ExtractedSection, RfcReference, SourceDocument

# Product codes: 1-2 letters, 4-6 digits, dash, 2-4 digits (e.g. M79060-001).
_PRODUCT_CODE_RE = re.compile(r"[A-Z]{1,2}[0-9]{4,6}-[0-9]{2,4}")
# Product-family code: just the base letters+digits without the revision suffix.
_PRODUCT_FAMILY_RE = re.compile(r"[A-Z]{1,2}[0-9]{4,6}")

# RFC workbook keyword used to classify .xlsx files.
_RFC_KEYWORDS = ("rfc",)

# Ordered by precedence: rfc_knowledge > debug_learning > hld > product_overview.
# Short acronyms (FA, RCA) are matched with word boundaries to avoid over-match.
_DEBUG_LEARNING_KEYWORDS = (
    "debug",
    "support",
    "key learning",
    "learning",
    "failure",
    "troubleshooting",
    "troubleshoot",
    "lesson",
    "known issue",
)
_DEBUG_LEARNING_WORD_KEYWORDS = ("fa", "rca")
_HLD_KEYWORDS = ("hld", "high level design", "high-level design", "architecture")
_PRODUCT_OVERVIEW_KEYWORDS = (
    "card",
    "product",
    "overview",
    "board",
    "module",
    "datasheet",
    "user guide",
    "manual",
    "spec",
)


def extract_product_code(filename: str) -> str | None:
    """Return the product code for ``filename``.

    Prefers the first ``_``/space-delimited token when it is itself a product
    code, then falls back to the first product-code match anywhere.
    """
    stem = os.path.splitext(os.path.basename(filename))[0]
    first_token = re.split(r"[_\s]+", stem, maxsplit=1)[0]
    token_match = _PRODUCT_CODE_RE.fullmatch(first_token)
    if token_match:
        return token_match.group(0)
    anywhere = _PRODUCT_CODE_RE.search(stem)
    return anywhere.group(0) if anywhere else None


def extract_product_family_code(filename: str) -> str | None:
    """Return the base product-family code (no revision suffix) for ``filename``.

    E.g. ``N32828-201_RFC.xlsx`` -> ``N32828``, ``N32828_RFC.xlsx`` -> ``N32828``.
    A revisioned product code's family is its base; if no revisioned code is
    found the family is derived directly from the bare family pattern.
    """
    stem = os.path.splitext(os.path.basename(filename))[0]
    # Prefer the base part of a full revisioned product code.
    full = _PRODUCT_CODE_RE.search(stem)
    if full:
        base = full.group(0).split("-")[0]
        m = _PRODUCT_FAMILY_RE.fullmatch(base)
        return m.group(0) if m else None
    # Fall back to a bare family code (letters+digits, no dash).
    first_token = re.split(r"[_\s]+", stem, maxsplit=1)[0]
    m = _PRODUCT_FAMILY_RE.fullmatch(first_token)
    if m:
        return m.group(0)
    anywhere = _PRODUCT_FAMILY_RE.search(stem)
    return anywhere.group(0) if anywhere else None


def detect_category(filename: str) -> DocumentCategory:
    """Classify a document by filename keywords (case-insensitive)."""
    name = os.path.basename(filename).lower()
    ext = os.path.splitext(name)[1]
    # XLSX files containing 'rfc' are RFC workbooks.
    if ext == ".xlsx" and any(kw in name for kw in _RFC_KEYWORDS):
        return "rfc_knowledge"
    # Tokenize on non-alphanumerics so `_fa_`/`_rca_` acronyms are found
    # (underscores are word chars, so \b alone would miss them).
    tokens = set(re.split(r"[^a-z0-9]+", name))
    if any(kw in name for kw in _DEBUG_LEARNING_KEYWORDS):
        return "debug_learning"
    if any(kw in tokens for kw in _DEBUG_LEARNING_WORD_KEYWORDS):
        return "debug_learning"
    if any(kw in name for kw in _HLD_KEYWORDS):
        return "hld"
    if any(kw in name for kw in _PRODUCT_OVERVIEW_KEYWORDS):
        return "product_overview"
    return "uncategorized"


def make_doc_id(product_code: str | None, filename: str) -> str:
    basis = f"{product_code or ''}|{os.path.basename(filename)}"
    return hashlib.sha1(basis.encode("utf-8")).hexdigest()[:12]


def scan_source_documents(
    source_dirs: list[str] | None = None,
    globs: list[str] | None = None,
) -> list[SourceDocument]:
    """Discover supported product documents across the configured source dirs."""
    dirs = source_dirs if source_dirs is not None else settings.PRODUCT_KNOWLEDGE_SOURCE_DIRS
    patterns = globs if globs is not None else settings.PRODUCT_KNOWLEDGE_SCAN_GLOBS
    seen: dict[str, SourceDocument] = {}
    for root in dirs:
        if not root or not os.path.isdir(root):
            continue
        for pattern in patterns:
            for path in glob.glob(os.path.join(root, "**", pattern), recursive=True):
                if not os.path.isfile(path):
                    continue
                # Skip Office lock/temp files (e.g. ~$20VMPDUPDB1_RFC_.xlsx).
                if os.path.basename(path).startswith("~$"):
                    continue
                doc = describe_document(path, source_root=root)
                # First discovery wins (stable across overlapping source dirs).
                seen.setdefault(doc.doc_id, doc)
    return sorted(seen.values(), key=lambda d: (d.product_code or "", d.filename))


def describe_document(path: str, source_root: str = "") -> SourceDocument:
    filename = os.path.basename(path)
    product_code = extract_product_code(filename)
    product_family_code = extract_product_family_code(filename)
    category = detect_category(filename)
    try:
        size_bytes = os.path.getsize(path)
    except OSError:
        size_bytes = 0
    doc = SourceDocument(
        doc_id=make_doc_id(product_code, filename),
        path=os.path.abspath(path),
        filename=filename,
        product_code=product_code,
        product_family_code=product_family_code,
        category=category,
        size_bytes=size_bytes,
        source_root=source_root,
    )
    if category == "rfc_knowledge" and product_family_code is None:
        doc.warnings.append(
            "RFC workbook filename contains no product code or product-family code; "
            "skipping ingestion."
        )
    elif product_code is None:
        doc.warnings.append("No product code found in filename.")
    return doc


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def _extract_pdf_lines(path: str) -> list[str]:
    from pypdf import PdfReader  # noqa: PLC0415 - optional/heavy import

    reader = PdfReader(path)
    lines: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        lines.extend(text.splitlines())
    return lines


def _extract_docx_blocks(path: str) -> list[tuple[bool, str]]:
    """Return (is_heading, text) blocks for a DOCX, in document order."""
    from docx import Document  # noqa: PLC0415 - optional/heavy import

    document = Document(path)
    blocks: list[tuple[bool, str]] = []
    for para in document.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style = (para.style.name if para.style else "") or ""
        is_heading = style.startswith("Heading") or style == "Title"
        blocks.append((is_heading, text))
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append((False, " | ".join(cells)))
    return blocks


# ---------------------------------------------------------------------------
# Sectioning
# ---------------------------------------------------------------------------

def _is_heading_line(line: str) -> bool:
    s = line.strip()
    if not s or len(s) > 90:
        return False
    if re.match(r"^\d+(\.\d+){0,3}[.\)]?\s+\S", s):
        return True
    letters = [c for c in s if c.isalpha()]
    if letters and len(s) <= 70 and s == s.upper():
        return True
    return False


def _chunk_text(text: str, max_chars: int) -> list[str]:
    """Split ``text`` into <= max_chars chunks on paragraph/line boundaries."""
    if len(text) <= max_chars:
        return [text]
    chunks: list[str] = []
    buf: list[str] = []
    size = 0
    for para in re.split(r"\n\s*\n", text):
        para = para.strip()
        if not para:
            continue
        if size + len(para) > max_chars and buf:
            chunks.append("\n\n".join(buf))
            buf, size = [], 0
        if len(para) > max_chars:
            # Hard-split an oversized paragraph.
            for i in range(0, len(para), max_chars):
                chunks.append(para[i : i + max_chars])
            continue
        buf.append(para)
        size += len(para)
    if buf:
        chunks.append("\n\n".join(buf))
    return chunks or [text[:max_chars]]


def _sections_from_headed_lines(
    lines: list[tuple[bool, str]],
) -> list[tuple[str | None, str]]:
    """Group (is_heading, text) pairs into (heading, body) sections."""
    sections: list[tuple[str | None, str]] = []
    heading: str | None = None
    body: list[str] = []

    def flush() -> None:
        text = "\n".join(body).strip()
        if text:
            sections.append((heading, text))

    for is_heading, text in lines:
        if is_heading:
            flush()
            heading = text
            body = []
        else:
            body.append(text)
    flush()
    return sections


# ---------------------------------------------------------------------------
# RFC XLSX extraction
# ---------------------------------------------------------------------------

# Column name variants accepted (case-insensitive, whitespace-normalised).
_RFC_COL_TEST = re.compile(r"failed.?test.?name|error.?name")
_RFC_COL_MSG = re.compile(r"error.?message|bin|issue|findings")
_RFC_COL_RFC = re.compile(r"^rfcs?(?:\s*\d+)?$")

# Separators for multiple RFC IDs inside one cell.
_RFC_ID_SEP = re.compile(r"[\n,;/\\\u2022\u25CF\u25E6]+")


def _normalize_col(name: object) -> str:
    return re.sub(r"\s+", " ", str(name or "").strip().lower())


def _find_header_row(ws) -> tuple[int, dict[str, list[int]]] | None:  # type: ignore[return]
    """Scan up to 20 rows for a header row; return (row_idx, col_map)."""
    for row_idx, row in enumerate(ws.iter_rows(max_row=20, values_only=True)):
        col_map: dict[str, list[int]] = {}
        for col_idx, cell in enumerate(row):
            norm = _normalize_col(cell)
            if _RFC_COL_TEST.search(norm) and "test" not in col_map:
                col_map["test"] = [col_idx]
            elif _RFC_COL_MSG.search(norm) and "msg" not in col_map:
                col_map["msg"] = [col_idx]
            elif _RFC_COL_RFC.fullmatch(norm):
                col_map.setdefault("rfc", []).append(col_idx)
        if "test" in col_map and col_map.get("rfc"):
            return row_idx, col_map
    return None


def _split_rfc_ids(cell_value: object) -> list[str]:
    """Split a cell value into individual RFC IDs, stripping empty parts."""
    raw = str(cell_value or "").strip()
    if not raw:
        return []
    return [p.strip() for p in _RFC_ID_SEP.split(raw) if p.strip()]


def _extract_xlsx_rfc_sections(
    path: str, product_code: str | None, product_family_code: str | None, doc_id: str
) -> tuple[str, list[ExtractedSection]]:
    """Parse an RFC workbook into one ExtractedSection per data row."""
    import hashlib as _hashlib

    from openpyxl import load_workbook  # noqa: PLC0415

    wb = load_workbook(path, read_only=True, data_only=True)
    all_text_parts: list[str] = []
    sections: list[ExtractedSection] = []
    order = 0

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        # Skip hidden sheets.
        if getattr(ws, "sheet_state", "visible") not in ("visible", None):
            continue
        result = _find_header_row(ws)
        if result is None:
            continue
        header_row_idx, col_map = result

        for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
            if row_idx <= header_row_idx:
                continue
            test_col = col_map["test"][0]
            msg_col = col_map.get("msg", [None])[0]
            test_val = row[test_col] if len(row) > test_col else None
            msg_val = row[msg_col] if (msg_col is not None and len(row) > msg_col) else None

            test_str = str(test_val or "").strip()
            msg_str = str(msg_val or "").strip() if msg_val is not None else ""
            rfc_entries: list[str] = []
            rfc_cols = col_map["rfc"]
            for ordinal, rfc_col in enumerate(rfc_cols, start=1):
                rfc_val = row[rfc_col] if len(row) > rfc_col else None
                for item in _split_rfc_ids(rfc_val):
                    if len(rfc_cols) > 1:
                        rfc_entries.append(f"RFC {ordinal}: {item}")
                    else:
                        rfc_entries.append(item)

            if not test_str or not rfc_entries:
                continue

            # Build stable human-readable text for the summarizer.
            parts = [f"failed_test_name: {test_str}"]
            if msg_str:
                parts.append(f"error_message_or_finding: {msg_str}")
            parts.append("rfc_entries:")
            parts.extend(f"- {entry}" for entry in rfc_entries)
            text = "\n".join(parts)
            heading = f"{sheet_name} — {test_str}"
            all_text_parts.append(text)

            sections.append(
                ExtractedSection(
                    section_id=f"{doc_id}-s{order:03d}",
                    doc_id=doc_id,
                    product_code=product_code,
                    category="rfc_knowledge",
                    heading=heading,
                    order=order,
                    text=text,
                )
            )
            order += 1

    wb.close()
    full_text = "\n".join(all_text_parts)
    if not full_text:
        raise ValueError("No RFC rows extracted from workbook.")
    content_hash = _hashlib.sha256(full_text.encode("utf-8")).hexdigest()[:16]
    return content_hash, sections


def parse_document(doc: SourceDocument) -> tuple[str, list[ExtractedSection]]:
    """Parse ``doc`` into bounded sections.

    Returns ``(content_hash, sections)``. Raises ``ValueError`` for unsupported
    extensions and when no text can be extracted.
    """
    # RFC workbooks with no product-family code cannot be usefully indexed.
    if doc.category == "rfc_knowledge" and doc.product_family_code is None:
        raise ValueError(
            f"{doc.filename}: RFC workbook filename contains no product code or "
            "product-family code; cannot ingest."
        )
    ext = os.path.splitext(doc.filename)[1].lower()
    if ext == ".xlsx":
        return _extract_xlsx_rfc_sections(
            doc.path, doc.product_code, doc.product_family_code, doc.doc_id
        )
    if ext == ".pdf":
        raw_lines = _extract_pdf_lines(doc.path)
        headed = [(_is_heading_line(ln), ln.strip()) for ln in raw_lines if ln.strip()]
    elif ext == ".docx":
        headed = _extract_docx_blocks(doc.path)
    else:
        raise ValueError(f"Unsupported document type: {ext}")

    full_text = "\n".join(text for _, text in headed).strip()
    if not full_text:
        raise ValueError("No extractable text in document.")
    content_hash = hashlib.sha256(full_text.encode("utf-8")).hexdigest()[:16]

    max_chars = settings.PRODUCT_KNOWLEDGE_SECTION_MAX_CHARS
    min_chars = settings.PRODUCT_KNOWLEDGE_SECTION_MIN_CHARS

    headed_sections = _sections_from_headed_lines(headed)
    has_headings = any(is_heading for is_heading, _ in headed)
    if not has_headings or not headed_sections:
        headed_sections = [(None, chunk) for chunk in _chunk_text(full_text, max_chars)]

    # Enforce max size, then merge undersized trailing sections forward.
    bounded: list[tuple[str | None, str]] = []
    for heading, text in headed_sections:
        for chunk in _chunk_text(text, max_chars):
            bounded.append((heading, chunk))

    # Merge undersized *headingless* chunks forward; never merge away a headed
    # section (that would lose document structure).
    merged: list[tuple[str | None, str]] = []
    for heading, text in bounded:
        if merged and heading is None and len(text) < min_chars:
            prev_heading, prev_text = merged[-1]
            if len(prev_text) + len(text) <= max_chars:
                merged[-1] = (prev_heading, f"{prev_text}\n{text}")
                continue
        merged.append((heading, text))

    sections: list[ExtractedSection] = []
    for order, (heading, text) in enumerate(merged):
        sections.append(
            ExtractedSection(
                section_id=f"{doc.doc_id}-s{order:03d}",
                doc_id=doc.doc_id,
                product_code=doc.product_code,
                category=doc.category,
                heading=heading,
                order=order,
                text=text,
            )
        )
    return content_hash, sections
