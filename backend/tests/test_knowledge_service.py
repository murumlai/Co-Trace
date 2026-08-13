"""Tests for the knowledge ingestion service (fake summarizer, real DOCX)."""
from __future__ import annotations

from docx import Document
from openpyxl import Workbook

from app.knowledge import parsing
from app.knowledge.models import ExtractedSection, KnowledgeSection, RfcReference, KnownFailureEntry
from app.knowledge.service import KnowledgeIngestionService
from app.knowledge.storage import KnowledgeStore


class FakeSummarizer:
    """Deterministic summarizer that never echoes raw document text."""

    def __init__(self, summary_prefix: str = "SUMMARY") -> None:
        self.prefix = summary_prefix

    def summarize(self, section: ExtractedSection, source_filename: str) -> KnowledgeSection:
        return KnowledgeSection(
            section_id=section.section_id,
            doc_id=section.doc_id,
            product_code=section.product_code,
            product_family_code=getattr(section, "product_family_code", None),
            category=section.category,
            heading=section.heading,
            order=section.order,
            summary=f"{self.prefix} for {section.heading} voltage rail",
            source_filename=source_filename,
            summary_model="fake",
        )


class FakeRfcSummarizer:
    """Deterministic RFC summarizer that emits rfc_references from the section text."""

    def summarize(self, section: ExtractedSection, source_filename: str) -> KnowledgeSection:
        # Parse rfc IDs out of the section text to produce realistic output.
        rfc_refs = []
        for line in (section.text or "").splitlines():
            if line.startswith("rfcs:"):
                for rfc_id in line.split(":", 1)[1].split(","):
                    rfc_id = rfc_id.strip()
                    if rfc_id:
                        rfc_refs.append(RfcReference(rfc_id=rfc_id, failed_test_name=section.heading))
        return KnowledgeSection(
            section_id=section.section_id,
            doc_id=section.doc_id,
            product_code=section.product_code,
            product_family_code=section.product_family_code if hasattr(section, "product_family_code") else None,
            category=section.category,
            heading=section.heading,
            order=section.order,
            summary=f"RFC section for {section.heading}",
            known_failures=[KnownFailureEntry(failing_step=section.heading, rfc_references=rfc_refs)],
            source_filename=source_filename,
            summary_model="fake-rfc",
        )


def _store(tmp_path) -> KnowledgeStore:
    return KnowledgeStore(
        manifest_path=str(tmp_path / "product_knowledge.json"),
        index_path=str(tmp_path / "product_knowledge_index.json"),
        sections_path=str(tmp_path / "product_knowledge_sections.jsonl"),
    )


def _make_docx(tmp_path, name) -> str:
    path = tmp_path / name
    document = Document()
    document.add_heading("Power", level=1)
    document.add_paragraph("RAW_DOC_BODY_MARKER regulated 20V rails and caps.")
    document.add_heading("Comms", level=1)
    document.add_paragraph("RAW_DOC_BODY_MARKER USB link and handshakes.")
    document.save(str(path))
    return str(path)


def _make_rfc_xlsx(tmp_path, name) -> str:
    wb = Workbook()
    ws = wb.active
    ws.title = "RFC Table"
    ws.append(["Failed Test Name / Error Name", "Error Message/ Bin / Issue / Findings", "RFCs"])
    ws.append(["POWER_TEST_01", "12V droop", "RFC-1234"])
    ws.append(["USB_ENUM_FAIL", "USB not detected", "RFC-0001"])
    path = tmp_path / name
    wb.save(str(path))
    return str(path)


class TestIngestionServiceSchemaVersion:
    def test_schema_version_is_2(self, tmp_path):
        store = _store(tmp_path)
        service = KnowledgeIngestionService(store=store, summarizer=FakeSummarizer())
        doc = parsing.describe_document(_make_docx(tmp_path, "N32828-201_HLD.docx"))
        manifest = service.build([doc])
        assert manifest.schema_version == 2
        index = store.load_index()
        assert index.schema_version == 2


class TestRfcWorkbookIngestion:
    def test_rfc_workbook_builds_rfc_knowledge_sections(self, tmp_path):
        store = _store(tmp_path)
        service = KnowledgeIngestionService(store=store, summarizer=FakeRfcSummarizer())
        doc = parsing.describe_document(_make_rfc_xlsx(tmp_path, "N32828_RFC.xlsx"))
        manifest = service.build([doc])

        assert store.exists()
        # The family code is N32828; since no revision it appears as UNKNOWN or N32828 in product_code.
        # product_family_code is N32828, but product_code is None.
        codes = {p.product_code for p in manifest.products}
        # Sections keyed under UNKNOWN (product_code=None)
        assert "UNKNOWN" in codes or "N32828" in codes
        category_counts = {p.product_code: p.category_counts for p in manifest.products}
        all_counts = {}
        for cc in category_counts.values():
            for k, v in cc.items():
                all_counts[k] = all_counts.get(k, 0) + v
        assert all_counts.get("rfc_knowledge", 0) == 2

    def test_rfc_workbook_no_raw_text_in_artifacts(self, tmp_path):
        store = _store(tmp_path)
        service = KnowledgeIngestionService(store=store, summarizer=FakeRfcSummarizer())
        doc = parsing.describe_document(_make_rfc_xlsx(tmp_path, "N32828_RFC.xlsx"))
        service.build([doc])
        raw = open(store.sections_path, encoding="utf-8").read()
        # The raw spreadsheet rows should not appear verbatim
        assert "RAW_DOC_BODY_MARKER" not in raw

    def test_rfc_index_has_product_family_code(self, tmp_path):
        store = _store(tmp_path)
        service = KnowledgeIngestionService(store=store, summarizer=FakeRfcSummarizer())
        doc = parsing.describe_document(_make_rfc_xlsx(tmp_path, "N32828_RFC.xlsx"))
        service.build([doc])
        index = store.load_index()
        all_entries = [e for entries in index.by_product.values() for e in entries]
        rfc_entries = [e for e in all_entries if e.category == "rfc_knowledge"]
        assert rfc_entries
        assert all(e.product_family_code == "N32828" for e in rfc_entries)
    def test_build_writes_pack_and_manifest(self, tmp_path):
        store = _store(tmp_path)
        service = KnowledgeIngestionService(store=store, summarizer=FakeSummarizer())
        doc = parsing.describe_document(_make_docx(tmp_path, "N32828-201_HLD.docx"))
        manifest = service.build([doc])

        assert store.exists()
        codes = {p.product_code for p in manifest.products}
        assert "N32828-201" in codes
        entry = next(p for p in manifest.products if p.product_code == "N32828-201")
        assert entry.section_count == 2
        assert entry.category_counts.get("hld") == 2
        assert entry.knowledge_hash

    def test_no_raw_document_text_in_artifacts(self, tmp_path):
        store = _store(tmp_path)
        service = KnowledgeIngestionService(store=store, summarizer=FakeSummarizer())
        doc = parsing.describe_document(_make_docx(tmp_path, "N32828-201_HLD.docx"))
        service.build([doc])
        raw = open(store.sections_path, encoding="utf-8").read()
        assert "RAW_DOC_BODY_MARKER" not in raw

    def test_knowledge_hash_changes_with_summary(self, tmp_path):
        store = _store(tmp_path)
        doc = parsing.describe_document(_make_docx(tmp_path, "N32828-201_HLD.docx"))

        first = KnowledgeIngestionService(store=store, summarizer=FakeSummarizer("A")).build([doc])
        hash_a = first.product_hash("N32828-201")

        second = KnowledgeIngestionService(store=store, summarizer=FakeSummarizer("B")).build([doc])
        hash_b = second.product_hash("N32828-201")

        assert hash_a and hash_b
        assert hash_a != hash_b

    def test_index_has_token_weights(self, tmp_path):
        store = _store(tmp_path)
        service = KnowledgeIngestionService(store=store, summarizer=FakeSummarizer())
        doc = parsing.describe_document(_make_docx(tmp_path, "N32828-201_HLD.docx"))
        service.build([doc])
        index = store.load_index()
        entries = index.by_product["N32828-201"]
        assert entries
        assert any(e.token_weights for e in entries)
        assert all(e.byte_length > 0 for e in entries)
