"""Tests for knowledge document discovery, categorization, and parsing."""
from __future__ import annotations

import pytest
from docx import Document

from app.knowledge import parsing


class TestProductCodeExtraction:
    @pytest.mark.parametrize(
        "filename, expected",
        [
            ("M79060-001_Debug Support Document Key Learning.pdf", "M79060-001"),
            ("M13983-700_Sedona USB Test Card.pdf", "M13983-700"),
            ("N32828-201_MPDU_20V_PDB1_HLD_Rev05.docx", "N32828-201"),
            ("K77469-400 overview.pdf", "K77469-400"),
            ("notes about M95113-001 board.pdf", "M95113-001"),
            ("random_notes.pdf", None),
        ],
    )
    def test_extract_product_code(self, filename, expected):
        assert parsing.extract_product_code(filename) == expected


class TestCategoryDetection:
    @pytest.mark.parametrize(
        "filename, expected",
        [
            ("M79060-001_Debug Support Document Key Learning.pdf", "debug_learning"),
            ("M13983-700_Sedona USB Test Card.pdf", "product_overview"),
            ("N32828-201_MPDU_20V_PDB1_HLD_Rev05.docx", "hld"),
            ("X_Troubleshooting_guide.pdf", "debug_learning"),
            ("Y_RCA_report.docx", "debug_learning"),
            ("Z_FA_summary.pdf", "debug_learning"),
            ("W_architecture_spec.docx", "hld"),
            ("V_datasheet.pdf", "product_overview"),
            ("plain_document.pdf", "uncategorized"),
        ],
    )
    def test_detect_category(self, filename, expected):
        assert parsing.detect_category(filename) == expected

    def test_short_acronyms_use_word_boundaries(self):
        # "sofa" contains "fa" but must not be mis-classified as debug_learning.
        assert parsing.detect_category("sofa_layout.pdf") == "uncategorized"


class TestSectioningHelpers:
    def test_chunk_text_respects_max(self):
        text = "\n\n".join(f"paragraph number {i} " * 10 for i in range(20))
        chunks = parsing._chunk_text(text, 200)
        assert len(chunks) > 1
        assert all(len(c) <= 400 for c in chunks)  # allow small paragraph overrun

    def test_sections_from_headed_lines_groups_bodies(self):
        lines = [
            (True, "Section One"),
            (False, "body a"),
            (False, "body b"),
            (True, "Section Two"),
            (False, "body c"),
        ]
        sections = parsing._sections_from_headed_lines(lines)
        assert [h for h, _ in sections] == ["Section One", "Section Two"]
        assert sections[0][1] == "body a\nbody b"

    def test_is_heading_line(self):
        assert parsing._is_heading_line("1.2 Power Subsystem")
        assert parsing._is_heading_line("OVERVIEW")
        assert not parsing._is_heading_line("this is a normal sentence of body text")


class TestParseDocx:
    def _make_docx(self, tmp_path, name, with_headings=True):
        path = tmp_path / name
        document = Document()
        if with_headings:
            document.add_heading("Architecture", level=1)
            document.add_paragraph("The MPDU provides regulated 20V rails.")
            document.add_heading("Interfaces", level=1)
            document.add_paragraph("USB and PDB1 connectors carry test signals.")
        else:
            for i in range(50):
                document.add_paragraph(f"UNIQUE_MARKER paragraph {i} with filler content here.")
        document.save(str(path))
        return str(path)

    def test_parses_headed_sections(self, tmp_path):
        path = self._make_docx(tmp_path, "N32828-201_HLD.docx")
        doc = parsing.describe_document(path)
        content_hash, sections = parsing.parse_document(doc)
        assert content_hash
        assert len(sections) == 2
        assert sections[0].heading == "Architecture"
        assert sections[0].product_code == "N32828-201"
        assert sections[0].category == "hld"
        assert all(s.section_id.startswith(doc.doc_id) for s in sections)

    def test_content_hash_is_stable(self, tmp_path):
        path = self._make_docx(tmp_path, "N32828-201_HLD.docx")
        doc = parsing.describe_document(path)
        first, _ = parsing.parse_document(doc)
        second, _ = parsing.parse_document(doc)
        assert first == second

    def test_heading_fallback_chunks_when_no_headings(self, tmp_path, monkeypatch):
        from app.config import settings

        monkeypatch.setattr(settings, "PRODUCT_KNOWLEDGE_SECTION_MAX_CHARS", 300)
        monkeypatch.setattr(settings, "PRODUCT_KNOWLEDGE_SECTION_MIN_CHARS", 10)
        path = self._make_docx(tmp_path, "M13983-700_card.docx", with_headings=False)
        doc = parsing.describe_document(path)
        _, sections = parsing.parse_document(doc)
        assert len(sections) > 1
        assert all(s.heading is None for s in sections)

    def test_unsupported_extension_raises(self, tmp_path):
        path = tmp_path / "x.txt"
        path.write_text("hello")
        doc = parsing.describe_document(str(path))
        with pytest.raises(ValueError):
            parsing.parse_document(doc)
